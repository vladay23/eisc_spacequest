# docx_generator.py
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

# ─── Вспомогательные функции (ГОСТ-оформление) ───────────────────────────────

def _setup_document(doc):
    """Настройка полей, шрифтов и интервалов по ГОСТ 2.105-95"""
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def _add_approval_block(doc):
    """Блок «УТВЕРЖДАЮ» в правом верхнем углу"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("УТВЕРЖДАЮ")
    run.bold = True
    run.font.size = Pt(12)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(6)
    p2.add_run("Заместитель начальника\nиспытательной лаборатории АО «SpaceQuest»")
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_after = Pt(12)
    p3.add_run("_________________   И.О. Фамилия\n«____» __________ 20__ г.")

def _add_centered_title(doc, title, number=""):
    """Заголовок документа по центру"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    
    if number:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run(f"№ {number}")
        run2.font.size = Pt(12)

def _add_section_header(doc, text, num=None):
    """Заголовок раздела с отступами по ГОСТ"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    prefix = f"{num}. " if num else ""
    run = p.add_run(f"{prefix}{text}")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)

def _add_formal_table(doc, headers, rows, col_widths=None):
    """Таблица в формальном стиле с сеткой и нумерацией первой колонки"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        
    for r_idx, r_data in enumerate(rows, start=1):
        table.add_row()
        data = list(r_data)
        # Если данных на 1 меньше заголовков, автоматически ставим номер строки
        if len(data) == len(headers) - 1:
            data.insert(0, str(r_idx))
            
        for c_idx, val in enumerate(data):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val) if val is not None else "—")
            run.font.size = Pt(10)
            
    if col_widths:
        for i, width in enumerate(col_widths[:len(table.columns)]):
            table.columns[i].width = Cm(width)
            
    return table

def _add_signature_block(doc, roles, date_issued=None):
    """Блок подписей в конце документа по ГОСТ"""
    doc.add_paragraph()
    if date_issued:
        p_date = doc.add_paragraph()
        p_date.paragraph_format.space_after = Pt(12)
        p_date.add_run(f"Дата выдачи документа: {date_issued}")
        
    for role in roles:
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_after = Pt(2)
        p_role.add_run(f"{role}:")
        
        p_sign = doc.add_paragraph()
        p_sign.paragraph_format.space_after = Pt(12)
        p_sign.add_run("_________________________   (_________________________)")
        
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(6)
        p_sub.add_run("      (подпись)                          (Фамилия И.О.)")
        p_sub.runs[0].font.size = Pt(9)
        p_sub.runs[0].italic = True

# ─── Генератор: Сводный отчёт по изделию (строго по main.py) ─────────────────

def generate_product_summary(product, components, mass_logs, fw_logs, test_runs, template_path=None):
    """Генерирует сводный отчёт по изделию в соответствии с ГОСТ 2.105-95"""
    doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    _setup_document(doc)
    
    _add_approval_block(doc)
    _add_centered_title(doc, "СВОДНЫЙ ОТЧЁТ ПО ИЗДЕЛИЮ", f"СВ-{product.serial_number}/{datetime.now().strftime('%y')}")
    
    # 1. Общие сведения
    _add_section_header(doc, "Общие сведения", num=1)
    status_map = {"design": "Разработка", "production": "Производство", "testing": "Испытания", "accepted": "Принято"}
    status_rus = status_map.get(product.status, product.status)
    
    _add_formal_table(doc, ["Параметр", "Значение"], [
        ("Наименование изделия", product.name or "—"),
        ("Заводской (серийный) номер", product.serial_number or "—"),
        ("Текущий статус жизненного цикла", status_rus),
        ("Масса (последнее измерение)", f"{product.mass_kg:.3f} кг" if product.mass_kg else "Не определена"),
        ("Соответствие требованиям по массе", "Соответствует" if product.within_tolerance else "Не соответствует"),
        ("Дата формирования отчёта", datetime.now().strftime('%d.%m.%Y')),
    ], col_widths=[8, 10])
    
    # 2. Состав и комплектация
    if components:
        _add_section_header(doc, "Состав и комплектация", num=2)
        rows = []
        for c in components:
            rows.append((c.name or "—", c.controller_type or "—", c.hardware_revision or "—", c.firmware_version or "Не установлено"))
        _add_formal_table(doc, ["№ п/п", "Наименование компонента", "Тип контроллера", "Ревизия ПО", "Примечание"], rows, col_widths=[1.5, 5, 4, 4, 3.5])
    
    # 3. Результаты контроля массы и центра масс
    if mass_logs:
        _add_section_header(doc, "Результаты контроля массы и центра масс", num=3)
        rows = []
        for log in mass_logs[:5]:
            cg_str = f"X={log.cg_x:.3f}; Y={log.cg_y:.3f}; Z={log.cg_z:.3f} м" if log.cg_x is not None else "Не определены"
            rows.append((
                log.checked_at.strftime('%d.%m.%Y') if log.checked_at else "—",
                f"{log.mass_kg:.3f}" if log.mass_kg else "—",
                "кг",
                cg_str,
                "±1.5 кг / ±0.05 м",
                "Соответствует" if log.within_tolerance else "Не соответствует"
            ))
        _add_formal_table(doc, ["№ п/п", "Дата контроля", "Фактическая масса", "Ед. изм.", "Положение ЦМ", "Требования НД", "Статус"], rows, col_widths=[1.5, 3, 2.5, 1.5, 4, 3.5, 2])
    
    # 4. Результаты испытаний
    if test_runs:
        _add_section_header(doc, "Результаты испытаний", num=4)
        rows = []
        for t in test_runs[:5]:
            scenario_name = t.scenario.name if hasattr(t, 'scenario') and t.scenario else "Не указан"
            status = "Пройдено" if t.status == "passed" else "Не пройдено"
            report = "Сформирован" if t.report_path else "Отсутствует"
            rows.append((t.started_at.strftime('%d.%m.%Y') if t.started_at else "—", scenario_name, status, report))
        _add_formal_table(doc, ["№ п/п", "Дата проведения", "Сценарий испытаний", "Результат", "Акт испытаний"], rows, col_widths=[1.5, 3, 4, 3, 3.5])
    
    # 5. Проверка совместимости бортового ПО
    if fw_logs:
        _add_section_header(doc, "Проверка совместимости бортового программного обеспечения", num=5)
        rows = []
        for log in fw_logs[:5]:
            comp_name = f"Компонент ID {log.component_id}"
            # Пытаемся найти имя компонента по ID из переданного списка
            comp_match = next((c for c in components if c.id == log.component_id), None)
            if comp_match:
                comp_name = comp_match.name
                
            status = "Совместимо" if log.compatible else "Несовместимо"
            rows.append((comp_name, log.checked_at.strftime('%d.%m.%Y') if log.checked_at else "—", status, "Матрица совместимости АО «SpaceQuest»"))
        _add_formal_table(doc, ["№ п/п", "Наименование компонента", "Дата проверки", "Результат", "Основание"], rows, col_widths=[1.5, 5, 3, 3, 4])
    
    # 6. Заключение и выводы
    _add_section_header(doc, "Заключение и выводы", num=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(1.25)
    
    issues = []
    if not product.within_tolerance:
        issues.append("выявлено отклонение массо-инерционных характеристик")
    if any(not f.compatible for f in fw_logs):
        issues.append("обнаружена несовместимость версий бортового ПО")
    if any(t.status == "failed" for t in test_runs):
        issues.append("зафиксированы отказы при проведении испытаний")
        
    if not issues:
        text = "По результатам комплексного анализа установлено, что изделие соответствует требованиям технической документации и готово к переходу на следующий этап жизненного цикла. Замечания отсутствуют."
    else:
        text = f"По результатам анализа выявлены следующие замечания: {', '.join(issues)}. Изделие требует проведения доработок и повторного контроля перед продолжением работ."
        
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    
    # Подписи
    _add_signature_block(doc, [
        "Сводный отчёт составил",
        "Проверил (ведущий инженер)",
        "Утвердил (начальник отдела качества)"
    ], date_issued=datetime.now().strftime('%d.%m.%Y'))
    
    # Сохранение
    filename = f"summary_{product.serial_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs("reports", exist_ok=True)
    filepath = os.path.join("reports", filename)
    doc.save(filepath)
    return filepath

# ─── Остальные генераторы (синхронизированы по стилю и сигнатурам main.py) ──

def generate_test_act(product, test_run, scenario, telemetry_points, template_path=None):
    doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    _setup_document(doc)
    _add_approval_block(doc)
    _add_centered_title(doc, "АКТ ИСПЫТАНИЙ", f"{test_run.id}/{datetime.now().strftime('%y')}")
    
    _add_section_header(doc, "Общие сведения", num=1)
    _add_formal_table(doc, ["Параметр", "Значение"], [
        ("Изделие", product.name),
        ("Серийный номер", product.serial_number),
        ("Сценарий испытаний", scenario.name),
        ("Идентификатор стенда", f"BENCH-{test_run.id:04d}"),
        ("Начало испытаний", test_run.started_at.strftime('%d.%m.%Y %H:%M:%S') if test_run.started_at else "—"),
        ("Окончание испытаний", test_run.finished_at.strftime('%d.%m.%Y %H:%M:%S') if test_run.finished_at else "—"),
        ("Количество точек телеметрии", len(telemetry_points)),
    ])
    
    _add_section_header(doc, "Допуски сценария", num=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run("Параметры допустимых зон (согласно ТД):")
    for param, limits in scenario.limits.items():
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        if isinstance(limits, dict):
            parts = [f"{k}: {v}" for k, v in limits.items()]
            p2.add_run(f"• {param}: {', '.join(parts)}")
        else:
            p2.add_run(f"• {param}: {limits}")
            
    _add_section_header(doc, "Результаты телеметрии", num=3)
    headers = ["№ п/п", "Параметр", "Фактическое значение", "Ед. изм.", "Статус"]
    rows = []
    for idx, (param, value, within) in enumerate(telemetry_points, 1):
        rows.append((param, f"{value:.3f}", "усл. ед.", "В пределах допуска" if within else "Отклонение"))
    table = _add_formal_table(doc, headers, rows, col_widths=[1.5, 4, 4, 2, 3.5])
    
    for row in table.rows[1:]:
        if "Отклонение" in row.cells[4].text:
            for p in row.cells[4].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(200, 0, 0)
                    run.bold = True
                    
    _add_section_header(doc, "Статистика испытаний", num=4)
    total = len(telemetry_points)
    passed = sum(1 for _, _, w in telemetry_points if w)
    _add_formal_table(doc, ["Общее число точек", "Соответствуют допуску", "Выявлены отклонения"], [(total, passed, total-passed)])
    
    _add_section_header(doc, "Заключение", num=5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Cm(1.25)
    if passed == total:
        run = p.add_run("Испытания проведены в полном объёме. Изделие соответствует требованиям технического задания. Допусков не выявлено.")
        run.font.color.rgb = RGBColor(0, 100, 0)
    else:
        run = p.add_run(f"Испытания не пройдены. Выявлено {total-passed} отклонений. Требуется проведение доработок и повторных испытаний.")
        run.font.color.rgb = RGBColor(180, 0, 0)
    run.bold = True
    
    _add_signature_block(doc, ["Инженер-испытатель", "Начальник испытательной лаборатории", "Представитель ОТК"], date_issued=datetime.now().strftime('%d.%m.%Y'))
    
    filename = f"test_act_{product.serial_number}_{test_run.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs("reports", exist_ok=True)
    filename = f"test_act_{product.serial_number}_{test_run.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join("reports", filename)
    doc.save(filepath)      # <-- ЭТА СТРОКА ОТСУТСТВОВАЛА
    return os.path.join("reports", filename)

def generate_mass_report(product, log, template_path=None):
    doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    _setup_document(doc)
    _add_approval_block(doc)
    _add_centered_title(doc, "ПРОТОКОЛ МАССОВО-ИНЕРЦИОННЫХ ХАРАКТЕРИСТИК", f"МИ-{product.serial_number}/{datetime.now().strftime('%y')}")
    
    _add_section_header(doc, "Сведения об изделии", num=1)
    _add_formal_table(doc, ["Параметр", "Значение"], [
        ("Наименование изделия", product.name),
        ("Серийный номер", product.serial_number),
        ("Статус", product.status),
        ("Дата проведения измерений", datetime.now().strftime('%d.%m.%Y')),
    ])
    
    _add_section_header(doc, "Результаты измерений", num=2)
    rows = [
        ("Масса изделия", f"{log.mass_kg:.3f}", "кг", "30.0 ± 1.5", "Соответствует" if log.within_tolerance else "Не соответствует"),
        ("Центр масс (X)", f"{log.cg_x:.3f}" if log.cg_x else "—", "м", "± 0.05", "Соответствует" if abs(log.cg_x or 0) < 0.05 else "Не соответствует"),
        ("Центр масс (Y)", f"{log.cg_y:.3f}" if log.cg_y else "—", "м", "± 0.05", "Соответствует" if abs(log.cg_y or 0) < 0.05 else "Не соответствует"),
        ("Центр масс (Z)", f"{log.cg_z:.3f}" if log.cg_z else "—", "м", "± 0.05", "Соответствует" if abs(log.cg_z or 0) < 0.05 else "Не соответствует"),
    ]
    table = _add_formal_table(doc, ["Наименование параметра", "Фактическое значение", "Ед. изм.", "Требования НД", "Статус"], rows, col_widths=[5, 3, 2, 4, 4])
    for row in table.rows[1:]:
        if "Не соответствует" in row.cells[4].text:
            for p in row.cells[4].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(200, 0, 0)
                    run.bold = True
                    
    _add_section_header(doc, "Заключение", num=3)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    if log.within_tolerance:
        run = p.add_run("Массо-инерционные характеристики изделия находятся в пределах допустимых значений, установленных конструкторской документацией.")
    else:
        run = p.add_run("Выявлены отклонения массовых характеристик. Требуется корректировка конструкции или перерасчёт компоновки.")
    run.bold = True
    
    _add_signature_block(doc, ["Инженер-конструктор", "Начальник КБ", "Представитель ОТК"], date_issued=datetime.now().strftime('%d.%m.%Y'))
    
    filename = f"mass_report_{product.serial_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs("reports", exist_ok=True)
    filepath = os.path.join("reports", filename)
    doc.save(filepath)
    return filepath

def generate_firmware_passport(product, component, log, template_path=None):
    doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    _setup_document(doc)
    _add_approval_block(doc)
    _add_centered_title(doc, "ПАСПОРТ СОВМЕСТИМОСТИ ПРОШИВКИ", f"ПО-{component.id}/{datetime.now().strftime('%y')}")
    
    _add_section_header(doc, "Сведения о компоненте", num=1)
    _add_formal_table(doc, ["Параметр", "Значение"], [
        ("Изделие", f"{product.name} (С/Н: {product.serial_number})"),
        ("Компонент", component.name),
        ("Тип контроллера", component.controller_type),
        ("Ревизия аппаратного обеспечения", component.hardware_revision),
    ])
    
    _add_section_header(doc, "Сведения о программном обеспечении", num=2)
    _add_formal_table(doc, ["Параметр", "Значение"], [
        ("Версия прошивки", component.firmware_version or "—"),
        ("Хэш-сумма (SHA-256)", (component.firmware_hash or "")[:32] + "..."),
        ("Дата проверки", log.checked_at.strftime('%d.%m.%Y %H:%M') if log.checked_at else "—"),
    ])
    
    _add_section_header(doc, "Результат проверки совместимости", num=3)
    table = _add_formal_table(doc, ["Критерий проверки", "Результат"], [
        ("Соответствие матрице совместимости АО «SpaceQuest»", "Совместимо" if log.compatible else "Несовместимо")
    ])
    status_cell = table.cell(1, 1)
    for p in status_cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0 if log.compatible else 180, 0)
            
    _add_section_header(doc, "Заключение", num=4)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    if log.compatible:
        run = p.add_run("Программное обеспечение допускается к установке и эксплуатации на указанном аппаратном обеспечении.")
    else:
        run = p.add_run("Установка данной версии прошивки не рекомендуется. Требуется обновление ПО или аппаратной ревизии компонента.")
        run.font.color.rgb = RGBColor(200, 100, 0)
    run.bold = True
    
    _add_signature_block(doc, ["Инженер ПО", "Руководитель направления", "Представитель службы качества"], date_issued=datetime.now().strftime('%d.%m.%Y'))
    
    filename = f"firmware_passport_{component.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs("reports", exist_ok=True)
    filepath = os.path.join("reports", filename)
    doc.save(filepath)
    return filepath

def generate_incidents_report(firmware_incidents, mass_incidents, test_failures, template_path=None):
    """Генерирует журнал инцидентов (несовместимости, отклонения по массе, проваленные тесты)"""
    doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    _setup_document(doc)
    
    _add_approval_block(doc)
    _add_centered_title(doc, "ЖУРНАЛ ИНЦИДЕНТОВ И ОТКЛОНЕНИЙ", f"ЖИ-{datetime.now().strftime('%Y%m%d')}")
    
    # 1. Несовместимость прошивок
    _add_section_header(doc, "Несовместимость прошивок с компонентами", num=1)
    if firmware_incidents:
        rows = []
        for log in firmware_incidents:
            rows.append((
                log.product.serial_number,
                log.product.name,
                f"{log.component.name} ({log.component.controller_type})",
                log.component.firmware_version or "—",
                log.component.hardware_revision,
                log.checked_at.strftime('%d.%m.%Y %H:%M') if log.checked_at else "—"
            ))
        _add_formal_table(doc, ["№ п/п", "С/Н изделия", "Наименование", "Компонент", "Версия ПО", "Ревизия железа", "Дата проверки"], rows, col_widths=[1.5, 3, 4, 4, 2.5, 2.5, 3])
    else:
        p = doc.add_paragraph()
        p.add_run("Несовместимых конфигураций не выявлено.")
        p.runs[0].italic = True
    
    # 2. Отклонения по массе
    _add_section_header(doc, "Отклонения по массе и центру масс", num=2)
    if mass_incidents:
        rows = []
        for log in mass_incidents:
            mass_str = f"{log.mass_kg:.3f} кг" if log.mass_kg else "—"
            cg_str = f"X={log.cg_x:.2f}, Y={log.cg_y:.2f}, Z={log.cg_z:.2f}" if log.cg_x else "—"
            rows.append((
                log.product.serial_number,
                log.product.name,
                mass_str,
                cg_str,
                log.material or "—",
                log.checked_at.strftime('%d.%m.%Y %H:%M') if log.checked_at else "—"
            ))
        _add_formal_table(doc, ["№ п/п", "С/Н изделия", "Наименование", "Масса", "Положение ЦМ", "Материал", "Дата проверки"], rows, col_widths=[1.5, 3, 4, 2.5, 4, 2.5, 3])
    else:
        p = doc.add_paragraph()
        p.add_run("Отклонений по массе и ЦМ не зафиксировано.")
        p.runs[0].italic = True
    
    # 3. Проваленные испытания
    _add_section_header(doc, "Испытания со статусом «Не пройдено»", num=3)
    if test_failures:
        rows = []
        for test in test_failures:
            rows.append((
                test.product.serial_number,
                test.product.name,
                test.scenario.name if test.scenario else "—",
                test.started_at.strftime('%d.%m.%Y %H:%M') if test.started_at else "—",
                test.finished_at.strftime('%d.%m.%Y %H:%M') if test.finished_at else "—",
                "Не пройдено"
            ))
        _add_formal_table(doc, ["№ п/п", "С/Н изделия", "Наименование", "Сценарий", "Начало", "Окончание", "Результат"], rows, col_widths=[1.5, 3, 4, 4, 3, 3, 2.5])
    else:
        p = doc.add_paragraph()
        p.add_run("Проваленных испытаний нет.")
        p.runs[0].italic = True
    
    # Заключение
    _add_section_header(doc, "Заключение", num=4)
    total_issues = len(firmware_incidents) + len(mass_incidents) + len(test_failures)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    if total_issues == 0:
        run = p.add_run("По состоянию на отчётную дату все изделия и компоненты соответствуют требованиям. Инциденты отсутствуют.")
        run.font.color.rgb = RGBColor(0, 100, 0)
    else:
        run = p.add_run(f"Выявлено {total_issues} инцидентов: {len(firmware_incidents)} по совместимости ПО, {len(mass_incidents)} по массе, {len(test_failures)} по испытаниям. Требуется проведение корректирующих мероприятий.")
        run.font.color.rgb = RGBColor(180, 100, 0)
    run.bold = True
    
    _add_signature_block(doc, [
        "Начальник отдела качества",
        "Ведущий инженер по надёжности",
        "Представитель технического управления"
    ], date_issued=datetime.now().strftime('%d.%m.%Y'))
    
    os.makedirs("reports", exist_ok=True)
    filename = f"incidents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join("reports", filename)
    doc.save(filepath)
    return filepath

def generate_analytics_report(
    total_products, in_tol, out_of_tol,
    passed_tests, failed_tests,
    compatibility_ok, compatibility_issues,
    readiness, readiness_text,
    filters: dict, template_path=None
):
    """Генерирует аналитический отчёт в формате Word (ГОСТ)"""
    doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    _setup_document(doc)
    
    _add_approval_block(doc)
    _add_centered_title(doc, "АНАЛИТИЧЕСКИЙ ОТЧЁТ ПО СТАТИСТИКЕ ИЗДЕЛИЙ", f"АН-{datetime.now().strftime('%Y%m%d')}")
    
    # --- Информация о фильтрах ---
    _add_section_header(doc, "Условия формирования отчёта", num=1)
    filter_rows = []
    if filters.get('date_from') and filters.get('date_to'):
        filter_rows.append(("Период", f"{filters['date_from']} – {filters['date_to']}"))
    elif filters.get('date_from'):
        filter_rows.append(("Период", f"с {filters['date_from']}"))
    elif filters.get('date_to'):
        filter_rows.append(("Период", f"по {filters['date_to']}"))
    if filters.get('status'):
        status_map = {"design": "Проектирование", "production": "Производство", "testing": "Испытания"}
        filter_rows.append(("Статус изделия", status_map.get(filters['status'], filters['status'])))
    if filters.get('issue_type'):
        issue_map = {"mass": "Отклонения по массе", "test": "Проваленные испытания", "firmware": "Несовместимость ПО"}
        filter_rows.append(("Тип отклонения", issue_map.get(filters['issue_type'], filters['issue_type'])))
    if not filter_rows:
        filter_rows.append(("Фильтры", "Не применялись"))
    _add_formal_table(doc, ["Параметр", "Значение"], filter_rows, col_widths=[5, 12])
    
    # --- Сводные карточки ---
    _add_section_header(doc, "Общая статистика", num=2)
    stats_rows = [
        ("Всего изделий в выборке", str(total_products)),
        ("Успешных испытаний", f"{passed_tests} (из {passed_tests + failed_tests} проведённых)"),
        ("Совместимость ПО", f"{compatibility_ok} совместимых компонентов из {compatibility_ok + compatibility_issues}"),
    ]
    _add_formal_table(doc, ["Показатель", "Значение"], stats_rows, col_widths=[6, 11])
    
    # --- Массо-инерционные характеристики ---
    _add_section_header(doc, "Массо-инерционные характеристики", num=3)
    total_mass = in_tol + out_of_tol
    percent_in = int(in_tol / total_mass * 100) if total_mass > 0 else 0
    percent_out = int(out_of_tol / total_mass * 100) if total_mass > 0 else 0
    mass_rows = [
        ("В пределах допуска", str(in_tol), f"{percent_in}%"),
        ("Вне допуска", str(out_of_tol), f"{percent_out}%"),
    ]
    _add_formal_table(doc, ["Параметр", "Количество", "Доля"], mass_rows, col_widths=[6, 3, 3])
    
    # --- Испытания ---
    _add_section_header(doc, "Результаты испытаний", num=4)
    total_tests = passed_tests + failed_tests
    percent_passed = int(passed_tests / total_tests * 100) if total_tests > 0 else 0
    percent_failed = int(failed_tests / total_tests * 100) if total_tests > 0 else 0
    test_rows = [
        ("Пройдено успешно", str(passed_tests), f"{percent_passed}%"),
        ("Не пройдено", str(failed_tests), f"{percent_failed}%"),
    ]
    _add_formal_table(doc, ["Параметр", "Количество", "Доля"], test_rows, col_widths=[6, 3, 3])
    
    # --- Совместимость ПО ---
    _add_section_header(doc, "Совместимость бортового ПО", num=5)
    total_compat = compatibility_ok + compatibility_issues
    percent_ok = int(compatibility_ok / total_compat * 100) if total_compat > 0 else 0
    percent_issues = int(compatibility_issues / total_compat * 100) if total_compat > 0 else 0
    comp_rows = [
        ("Совместимо", str(compatibility_ok), f"{percent_ok}%"),
        ("Несовместимо", str(compatibility_issues), f"{percent_issues}%"),
    ]
    _add_formal_table(doc, ["Параметр", "Количество", "Доля"], comp_rows, col_widths=[6, 3, 3])
    
    # --- Индекс готовности ---
    _add_section_header(doc, "Интегральная оценка готовности", num=6)
    readiness_grade = "Высокий" if readiness >= 80 else "Средний" if readiness >= 50 else "Низкий"
    total_checks = (in_tol + out_of_tol) + (passed_tests + failed_tests) + (compatibility_ok + compatibility_issues)
    passed_checks = in_tol + passed_tests + compatibility_ok
    ready_rows = [
        ("Индекс готовности", f"{readiness}%", readiness_grade),
        ("Успешных проверок", str(passed_checks), f"из {total_checks}"),
    ]
    _add_formal_table(doc, ["Показатель", "Значение", "Оценка"], ready_rows, col_widths=[5, 4, 4])
    
    # --- Заключение ---
    _add_section_header(doc, "Заключение", num=7)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(readiness_text)
    run.bold = True
    if readiness < 100:
        run.font.color.rgb = RGBColor(180, 100, 0)
    else:
        run.font.color.rgb = RGBColor(0, 100, 0)
    
    _add_signature_block(doc, [
        "Главный инженер",
        "Начальник отдела качества",
        "Аналитик"
    ], date_issued=datetime.now().strftime('%d.%m.%Y'))
    
    os.makedirs("reports", exist_ok=True)
    filename = f"analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join("reports", filename)
    doc.save(filepath)
    return filepath

